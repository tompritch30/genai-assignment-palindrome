"""Document loader for parsing .docx files."""

from pathlib import Path
from typing import BinaryIO

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from src.utils.logging_config import get_logger

logger = get_logger(__name__)


class InvalidFileError(Exception):
    """Raised when file is not a valid .docx file."""

    pass


class EmptyDocumentError(Exception):
    """Raised when document contains no text."""

    pass


class DocumentLoader:
    """Loader for extracting text from .docx files."""

    @staticmethod
    def load_from_file(file_path: str | Path) -> str:
        """Load and extract text from a .docx file.

        Args:
            file_path: Path to the .docx file

        Returns:
            Extracted text content from the document

        Raises:
            InvalidFileError: If file is not a valid .docx or doesn't exist
            EmptyDocumentError: If document contains no text
        """
        file_path = Path(file_path)

        # Validate file exists
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file extension
        if file_path.suffix.lower() != ".docx":
            logger.error(f"Invalid file type: {file_path.suffix}")
            raise InvalidFileError(
                f"Expected .docx file, got {file_path.suffix}. File: {file_path}"
            )

        try:
            # Load document
            doc = Document(file_path)
            logger.info(f"Successfully loaded document: {file_path}")

            # Extract text from all paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(text_parts).strip()

            # Check if document is empty
            if not text:
                logger.warning(f"Document appears to be empty: {file_path}")
                raise EmptyDocumentError(f"Document contains no text: {file_path}")

            logger.info(f"Extracted {len(text)} characters from document")
            return text

        except EmptyDocumentError:
            raise
        except PackageNotFoundError:
            logger.error(
                f"Invalid .docx file (corrupted or not a Word document): {file_path}"
            )
            raise InvalidFileError(
                f"File is not a valid .docx document (may be corrupted): {file_path}"
            )
        except Exception as e:
            logger.error(f"Unexpected error loading document {file_path}: {e}")
            raise InvalidFileError(f"Failed to load document: {e}")

    @staticmethod
    def load_from_bytes(file_bytes: bytes, filename: str = "document.docx") -> str:
        """Load and extract text from .docx file bytes.

        Args:
            file_bytes: Raw bytes of the .docx file
            filename: Original filename (for logging)

        Returns:
            Extracted text content from the document

        Raises:
            InvalidFileError: If bytes are not a valid .docx
            EmptyDocumentError: If document contains no text
        """
        try:
            # Create document from bytes
            from io import BytesIO

            doc = Document(BytesIO(file_bytes))
            logger.info(f"Successfully loaded document from bytes: {filename}")

            # Extract text from all paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(text_parts).strip()

            # Check if document is empty
            if not text:
                logger.warning(f"Document appears to be empty: {filename}")
                raise EmptyDocumentError(f"Document contains no text: {filename}")

            logger.info(f"Extracted {len(text)} characters from document")
            return text
        except EmptyDocumentError:
            # Let EmptyDocumentError propagate as-is
            raise
        except PackageNotFoundError:
            logger.error(
                f"Invalid .docx file (corrupted or not a Word document): {filename}"
            )
            raise InvalidFileError(
                f"File is not a valid .docx document (may be corrupted): {filename}"
            )
        except Exception as e:
            logger.error(f"Unexpected error loading document {filename}: {e}")
            raise InvalidFileError(f"Failed to load document: {e}")

    @staticmethod
    def load_from_stream(stream: BinaryIO, filename: str = "document.docx") -> str:
        """Load and extract text from a file-like stream.

        Args:
            stream: Binary file-like object (e.g., from Streamlit upload)
            filename: Original filename (for logging)

        Returns:
            Extracted text content from the document

        Raises:
            InvalidFileError: If stream is not a valid .docx
            EmptyDocumentError: If document contains no text
        """
        try:
            # Read stream and create document
            doc = Document(stream)
            logger.info(f"Successfully loaded document from stream: {filename}")

            # Extract text from all paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(text_parts).strip()

            # Check if document is empty
            if not text:
                logger.warning(f"Document appears to be empty: {filename}")
                raise EmptyDocumentError(f"Document contains no text: {filename}")

            logger.info(f"Extracted {len(text)} characters from document")
            return text

        except EmptyDocumentError:
            # Let EmptyDocumentError propagate as-is
            raise
        except PackageNotFoundError:
            logger.error(
                f"Invalid .docx file (corrupted or not a Word document): {filename}"
            )
            raise InvalidFileError(
                f"File is not a valid .docx document (may be corrupted): {filename}"
            )
        except Exception as e:
            logger.error(f"Unexpected error loading document {filename}: {e}")
            raise InvalidFileError(f"Failed to load document: {e}")
